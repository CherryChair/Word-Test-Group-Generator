from groups import input_numbers, create_task_division
from docx import Document
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT


def get_task_division(doc_info):
    task_division = create_task_division(doc_info)
    tasks_sorted = []
    for task in task_division:
        tasks_sorted.append(sorted(task_division[task]))
    return tasks_sorted


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def make_columns_bold(*columns):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def populate_word_doc(document, task_division):
    people_number = len(task_division)
    tasks_number = len(task_division[0])
    document.add_paragraph()
    table = document.add_table(people_number + 1, 2)
    table.cell(0, 0).text = "Nr"
    table.cell(0, 1).text = "Zadania"
    table.cell(0, 1).add_paragraph()
    small_table = table.cell(0, 1).add_table(2, tasks_number)
    small_table.style = "Table Grid"
    for task_number in range(0, tasks_number):
        small_table.cell(0, task_number).text = f"Zad.{task_number + 1}"
        small_table.cell(1, task_number).text = "p"
    table.cell(0, 1).add_paragraph("Punktacja: 0-8 ndst, 8,5p-12p dop, 12,5p – 17,5p dst, 18 p – 21,5p db , 22p-24p bdb")
    make_rows_bold(small_table.rows[0], small_table.rows[1])
    for person_number in range(1, people_number + 1):
        table.cell(person_number, 0).text = f"{person_number}"
        cell = table.cell(person_number, 1)
        cell.add_paragraph()
        paragraph = cell.paragraphs[1]
        run = paragraph.add_run()
        try:
            for task_number in range(1, tasks_number + 1):
                # cell.add_paragraph()
                # paragraph = cell.add_paragraph()
                run = paragraph.add_run()
                current_task = task_division[person_number - 1][task_number - 1]
                task_group_nums = current_task.split("_")
                group_num = task_group_nums[1]
                run.add_picture("zadania/" + group_num + f"/{task_number}" + ".png", width=Cm(17.5))
        except Exception:
            print("Niepoprawnie ustawiony folder z zadaniami. Popraw go i spróbuj ponownie.")
            input("Wciśnij enter, żeby wyjść. ")
            exit()
        cell.add_paragraph()
    make_rows_bold(table.rows[0])
    make_columns_bold(table.columns[0])
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Cm(1.0)
    for cell in table.columns[1].cells:
        cell.width = Cm(18.0)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def get_list_of_occurances(task_division):
    list_of_occurances = []
    for i in range(1, len(task_division) + 1):
        occurances_dict = {}
        for j in range(0, len(task_division[0]) + 1):
            occurances_dict[j] = 0
        list_of_occurances.append(occurances_dict)
    task_dict = {i+1: set(task_division[i]) for i in range(0, len(task_division))}
    for number in task_dict:
        for person in task_dict:
            same_elements = len(task_dict[number].intersection(task_dict[person]))
            list_of_occurances[number-1][same_elements] += 1
        list_of_occurances[number-1][len(task_dict[1])] -= 1
    return list_of_occurances


def get_table_of_occurances(list_of_occurances):
    person = 0
    titular_string = "    "
    for number in range(0, len(list_of_occurances[0])):
        if number in [2, 3, 4]:
            titular_string += f"{number} powtórzenia".center(15)
        elif number == 1:
            titular_string += "1 powtórzenie".center(15)
        else:
            titular_string += f"{number} powtórzeń".center(15)
    titular_string += "\n"
    occurances_string = ""
    for person_occurances in list_of_occurances:
        person += 1
        occurances_string += f"Nr {person} "
        if person < 10:
            occurances_string += " "
        for index in person_occurances:
            if index == 0:
                occurances_string += f"{person_occurances[index]:>7}"
            else:
                occurances_string += f"{person_occurances[index]:>15}"
        occurances_string += "\n"
    return titular_string + occurances_string


def create_word_doc(task_division, filename):
    word_doc = Document()
    populate_word_doc(word_doc, task_division)
    word_doc.save("sprawdziany/" + filename + ".docx")


def create_table_of_occurances(task_division):
    list_of_occurances = get_list_of_occurances(task_division)
    table_of_occurances = get_table_of_occurances(list_of_occurances)
    return table_of_occurances


if __name__ == "__main__":
    doc_info = input_numbers()
    task_division = get_task_division(doc_info)
    filename = input("Podaj nazwę pliku sprawdzianu: ")
    create_word_doc(task_division, filename)
    table_of_occurances = create_table_of_occurances(task_division)
    print("\n" + table_of_occurances)
    while True:
        user_choice = input("Wybierz opcję i zatwierdź enterem:\n1 - dla tych samych danych " +
                            "wylosuj inny rozkład zadań i zapisz w pliku o podanej wcześniej" +
                            " nazwie\n2 - to samo co 1, tylko zapisz w pliku o innej nazwie\n" +
                            "3 - wpisz nowe dane i wylosuj nowy plik\n0 - zapisz i wyjdź\nWybór opcji: ").strip()
        try:
            user_choice = int(user_choice)
        except Exception:
            continue
        if user_choice == 1:
            task_division = get_task_division(doc_info)
            create_word_doc(task_division, filename)
            table_of_occurances = create_table_of_occurances(task_division)
            print("\n" + table_of_occurances)
            continue
        if user_choice == 0:
            break
        if user_choice == 2:
            task_division = get_task_division(doc_info)
            filename = input("Podaj nazwę pliku sprawdzianu: ")
            create_word_doc(task_division, filename)
            table_of_occurances = create_table_of_occurances(task_division)
            print("\n" + table_of_occurances)
            continue
        if user_choice == 3:
            doc_info = input_numbers()
            task_division = get_task_division(doc_info)
            filename = input("Podaj nazwę pliku sprawdzianu: ")
            create_word_doc(task_division, filename)
            table_of_occurances = create_table_of_occurances(task_division)
            print("\n" + table_of_occurances)
            continue
