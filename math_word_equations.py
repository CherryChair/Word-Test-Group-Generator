from groups import input_numbers, create_task_division
from docx import Document
from docx.shared import Cm


def get_task_division():
    table = input_numbers()
    task_division = create_task_division(table)
    tasks_sorted = []
    for task in task_division:
        tasks_sorted.append(sorted(task_division[task]))
    return tasks_sorted


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def make_columns_bold(*columns):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def populate_word_doc(document, task_division):
    people_number = len(task_division)
    tasks_number = len(task_division[0])
    table = document.add_table(people_number + 1, 2)
    table.cell(0, 0).text = "Nr"
    table.cell(0, 1).text = "Zadania"
    for person_number in range(1, people_number + 1):
        table.cell(person_number, 0).text = f"{person_number}"
        cell = table.cell(person_number, 1)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        for task_number in range(1, tasks_number + 1):
            current_task = task_division[person_number - 1][task_number - 1]
            run.add_picture("zadania/" + current_task + ".png", width=Cm(16.0))
    make_rows_bold(table.rows[0])
    make_columns_bold(table.columns[0])


if __name__ == "__main__":
    task_division = get_task_division()
    word_doc = Document()
    populate_word_doc(word_doc, task_division)
    filename = input("Podaj nazwę pliku sprawdzianu: ")
    word_doc.save("sprawdziany/" + filename + ".docx")
