import random
import docx
import matplotlib.mathtext as mathtext
EXPRESSION=r'$(\frac{5 - \frac{1}{x}}{4})$'


free_people_1 = set(range(6, 39))
free_people_2 = set(range(6, 39))
free_people_3 = set(range(6, 39))
free_people_4 = set(range(6, 39))
free_people_5 = set(range(6, 39))

free_tasks = [free_people_1, free_people_2, free_people_3, free_people_4, free_people_5]

set1 = ["1.1", "2.1", "3.1", "4.1", "5.1"]
set2 = ["1.2", "2.2", "3.2", "4.2", "5.2"]
set3 = ["1.3", "2.3", "3.3", "4.3", "5.3"]
set4 = ["1.4", "2.4", "3.4", "4.4", "5.4"]
set5 = ["1.5", "2.5", "3.5", "4.5", "5.5"]


task_dict = {1: set1, 2: set2, 3: set3, 4: set4, 5: set5}

for i in range(6, 39):
    task_dict[i] = []

for j in range(1, 6):
    for i in range(1, 6):
        if j == 5:
            chosen_people = set(random.sample(free_tasks[i-1], 5))
            free_tasks[i-1].difference_update(chosen_people)
        else:
            chosen_people = set(random.sample(free_tasks[i-1], 7))
            free_tasks[i-1].difference_update(chosen_people)
        for person in chosen_people:
            task_dict[person].append(f"{j}.{i}")


for number in task_dict:
    task_dict[number] = set(task_dict[number])

for i in range(1, 39):
    exec(f"inter_{i} = " + "{0:0, 1:0, 2:0, 3:0, 4:0, 5:0}")

for number in task_dict:
    for person in task_dict:
        same_elements = len(task_dict[number].intersection(task_dict[person]))
        exec(f"inter_{number}[{same_elements}] += 1")
    exec(f"inter_{number}[5]-=1")


word_doc = docx.Document("Sprawozdanie z laboratorium z automatyki WAET.docx")
a = word_doc.paragraphs
new_paragraph = a[0]
word_doc.add_paragraph("aaaaaaa")
word_doc.save("apple.docx")
# word_doc.save("/home/michal/python/word_doc/Sprawozdanie z laboratorium z automatyki WAET.docx")


pass
